# backend/app.py
from __future__ import annotations
import os
import logging
from datetime import datetime, timedelta, timezone
from functools import wraps
from pathlib import Path
from urllib.parse import urlparse

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_bcrypt import Bcrypt
from werkzeug.utils import secure_filename

from dotenv import load_dotenv
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError
from sqlalchemy import func

# Google Login deps
from google.oauth2 import id_token
from google.auth.transport import requests as g_requests

# Local modules
import models
import database

# Lazy import heavy dependencies only when needed
# import fitz  # PyMuPDF - moved to function
# import docx - moved to function
# from PIL import Image - moved to function
# import pytesseract - moved to function
# from vector_store import add_document, query_similar - moved to function

# New deps for URL importing
import requests
from bs4 import BeautifulSoup

# ---- Load env early ----
load_dotenv(dotenv_path=Path(__file__).with_name(".env"))

# ---- App setup ----
app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "dev-secret")
# Log a short preview so you can verify the .env key is loaded on every restart
print("🔐 Loaded SECRET_KEY:", (app.config["SECRET_KEY"][:10] + "...") if app.config["SECRET_KEY"] else "MISSING")
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB
UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# CORS
CORS(app, resources={
    r"/*": {  # Changed from r"/api/*" to r"/*" to allow all routes including /health
        "origins": [
            "http://localhost:5173",
            "http://localhost:5174",  # Added port 5174
            "http://localhost:3000",
            "http://127.0.0.1:5173",
            "http://127.0.0.1:5174",  # Added port 5174
            "http://127.0.0.1:3000",
            "https://smartdocq-gfzj.onrender.com"
        ],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "supports_credentials": True,
        "expose_headers": ["Content-Type", "Authorization"]
    }
})

logging.basicConfig(level=logging.DEBUG)
bcrypt = Bcrypt(app)

# Add request logging middleware
@app.before_request
def log_request_info():
    app.logger.debug('Headers: %s', request.headers)
    app.logger.debug('Body: %s', request.get_data())
    app.logger.info(f'Request: {request.method} {request.path}')

# ---- DB: ensure tables exist ----
models.Base.metadata.create_all(bind=database.engine)

# ---- AI toggles / config ----
AI_ENABLED = os.getenv("AI_ENABLED", "true").lower() == "true"
TESTING_MODE_ON_RATE_LIMIT = os.getenv("TESTING_MODE_ON_RATE_LIMIT", "false").lower() == "true"

# Model name is best-effort; prefer the requested alternate model by default.
# This still respects the MODEL_NAME env var if you set it in .env.
MODEL_NAME = os.getenv("MODEL_NAME", "models/gemini-2.5-flash-lite")
try:
    TEMPERATURE = float(os.getenv("TEMPERATURE", "0.2"))
except Exception:
    TEMPERATURE = 0.2

# Auth token lifetime (hours). Make this long so restarts don’t log you out.
TOKEN_TTL_HOURS = int(os.getenv("TOKEN_TTL_HOURS", "720"))  # default 30 days

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Import google.generativeai only if enabled+key present
genai = None
if AI_ENABLED and GOOGLE_API_KEY:
    try:
        import google.generativeai as genai  # type: ignore
        genai.configure(api_key=GOOGLE_API_KEY)
        # Prefer the chosen model first, then fall back to other candidates.
        model_candidates = [
            "models/gemini-2.5-flash-lite",
            "models/gemini-2.5-flash",
            "models/gemini-flash-latest",
            "models/gemini-1.5-flash",
            "models/gemini-1.5",
            "models/gemini-2.5-pro",
            "models/gemini-pro-latest",
        ]
        try:
            available_models = [m.name for m in genai.list_models()]
            for m in model_candidates:
                if m in available_models:
                    MODEL_NAME = m
                    break
        except Exception:
            app.logger.debug("Could not list models; keeping configured/default MODEL_NAME")
        app.logger.info(f"AI is ENABLED. Using model: {MODEL_NAME}")
    except Exception as e:
        app.logger.error(f"Failed to init Google Generative AI SDK: {e}")
        genai = None
else:
    app.logger.info("AI is DISABLED (no calls will be made).")

# =========================
# Utilities
# =========================
def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.teardown_appcontext
def shutdown_session(exception=None):
    """Ensure SQLAlchemy sessions are cleaned up after each request/app context."""
    try:
        database.SessionLocal.remove()  # if SessionLocal is a scoped_session
    except Exception:
        pass

def jwt_encode(payload: dict, secret: str) -> str:
    import jwt  # PyJWT
    return jwt.encode(payload, secret, algorithm="HS256")

def jwt_decode(token: str, secret: str) -> dict:
    import jwt  # PyJWT
    return jwt.decode(token, secret, algorithms=["HS256"])

def _issue_token_for(email: str) -> str:
    return jwt_encode(
        {"email": email, "exp": datetime.now(timezone.utc) + timedelta(hours=TOKEN_TTL_HOURS)},
        app.config["SECRET_KEY"],
    )

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        if "Authorization" in request.headers:
            parts = request.headers["Authorization"].split(" ")
            if len(parts) == 2 and parts[0].lower() == "bearer":
                token = parts[1]

        if not token:
            return jsonify({"message": "Token is missing!"}), 401

        db: Session | None = None
        try:
            data = jwt_decode(token, app.config["SECRET_KEY"])
            email = data.get("email")
            role = data.get("role", "student")
            
            # For admin tokens, create a mock user object
            if role == "admin":
                class AdminUser:
                    def __init__(self):
                        self.email = email
                        self.role = "admin"
                        self.id = None
                        self.name = "Administrator"
                current_user = AdminUser()
            else:
                db = database.SessionLocal()
                current_user = db.query(models.User).filter_by(email=email).first()
                if not current_user:
                    return jsonify({"message": "User not found!"}), 401
        except Exception as e:
            app.logger.error(f"Token validation failed: {str(e)}")
            return jsonify({"message": "Token is invalid!", "error": str(e)}), 401
        finally:
            if db is not None:
                db.close()

        return f(current_user, *args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(current_user, *args, **kwargs):
        user_role = getattr(current_user, "role", "student")
        app.logger.info(f"Admin check - User role: {user_role}")
        if user_role != "admin":
            return jsonify({"message": "Admin access required!"}), 403
        return f(current_user, *args, **kwargs)
    return decorated

def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF; OCR each page if needed."""
    import fitz  # Lazy import
    from PIL import Image
    import pytesseract
    
    text = ""
    with fitz.open(filepath) as doc:
        for page in doc:
            page_text = page.get_text()
            if not page_text.strip():
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                page_text = pytesseract.image_to_string(img)
            text += page_text + "\n"
    return text

def extract_text_from_docx(filepath: str) -> str:
    import docx  # Lazy import
    d = docx.Document(filepath)
    return "\n".join(p.text for p in d.paragraphs)


def convert_docx_to_html(filepath: str) -> str:
    """Enhanced DOCX -> HTML conversion with better formatting."""
    try:
        import docx
        from docx.shared import Pt
        
        d = docx.Document(filepath)
        parts = ['<div class="docx-html" style="font-family: Arial, sans-serif; line-height: 1.6; color: #e5e7eb;">']
        
        for p in d.paragraphs:
            if not p.text.strip():
                continue
                
            text = (p.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            style_name = (p.style.name or "").lower() if hasattr(p, 'style') and p.style else ""
            
            # Handle headings
            if 'heading 1' in style_name or style_name.startswith('h1'):
                parts.append(f'<h1 style="color: #fff; font-size: 24px; margin: 16px 0 8px;">{text}</h1>')
            elif 'heading 2' in style_name or style_name.startswith('h2'):
                parts.append(f'<h2 style="color: #fff; font-size: 20px; margin: 14px 0 7px;">{text}</h2>')
            elif 'heading 3' in style_name or style_name.startswith('h3'):
                parts.append(f'<h3 style="color: #fff; font-size: 18px; margin: 12px 0 6px;">{text}</h3>')
            else:
                # Regular paragraph with run formatting
                formatted_text = ""
                for run in p.runs:
                    run_text = (run.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    style_attrs = []
                    
                    if run.bold:
                        style_attrs.append("font-weight: bold")
                    if run.italic:
                        style_attrs.append("font-style: italic")
                    if run.underline:
                        style_attrs.append("text-decoration: underline")
                    
                    if style_attrs:
                        formatted_text += f'<span style="{"; ".join(style_attrs)}">{run_text}</span>'
                    else:
                        formatted_text += run_text
                
                parts.append(f'<p style="margin: 8px 0;">{formatted_text or text}</p>')
        
        parts.append("</div>")
        return "\n".join(parts)
    except Exception as e:
        import traceback
        app.logger.error(f"DOCX->HTML conversion failed: {str(e)}\n{traceback.format_exc()}")
        return f'<div style="color: #fca5a5; padding: 12px;"><p>Failed to convert document to HTML.</p><pre>{str(e)}</pre></div>'

def extract_text_from_image(filepath: str) -> str:
    from PIL import Image
    import pytesseract
    img = Image.open(filepath).convert("RGB")
    return pytesseract.image_to_string(img)

def extract_text_from_url(url: str) -> str:
    """Fetch HTML and return readable text (basic boilerplate removal)."""
    resp = requests.get(url, timeout=15)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in text.splitlines()]
    chunks = [ln for ln in lines if ln]
    return "\n".join(chunks)

def safe_answer_mock(reason: str) -> dict:
    return {
        "answer": f"(mock) AI disabled or unavailable: {reason}",
        "_used_model": None,
        "_mode": "mock"
    }

# =========================
# Auth Routes
# =========================
@app.post("/api/register")
def register():
    db: Session = next(get_db())
    data = request.get_json() or {}
    try:
        email = (data.get("email") or "").strip().lower()
        password = data.get("password") or ""
        name = data.get("name") or None

        if not email or not password:
            return jsonify({"message": "Missing data"}), 400

        if db.query(models.User).filter(models.User.email == email).first():
            return jsonify({"message": "User already exists!"}), 409

        hashed_password = bcrypt.generate_password_hash(password).decode("utf-8")
        new_user = models.User(name=name, email=email, hashed_password=hashed_password)
        db.add(new_user)
        db.commit()
        db.refresh(new_user)
        return jsonify({"message": "Registered successfully!"}), 201
    except IntegrityError:
        db.rollback()
        return jsonify({"message": "User already exists!"}), 409
    except Exception as e:
        db.rollback()
        app.logger.exception("Error in /api/register")
        return jsonify({"message": "Server error during registration", "error": str(e)}), 500
    finally:
        db.close()

@app.post("/api/login")
def login():
    db: Session = next(get_db())
    data = request.get_json() or {}
    try:
        email = (data.get("email") or "").strip().lower()
        password = data.get("password") or ""
        
        # Debug logging
        app.logger.info(f"Login attempt - Email: {email}")
        app.logger.info(f"Request origin: {request.headers.get('Origin')}")
        app.logger.info(f"Request host: {request.headers.get('Host')}")
        
        # Check for fixed admin credentials (case-insensitive email)
        if email == ADMIN_EMAIL and password == ADMIN_PASSWORD:
            app.logger.info(f"Admin login successful for: {email}")
            # Return admin token with role
            token = jwt_encode(
                {"email": email, "role": "admin", "exp": datetime.now(timezone.utc) + timedelta(hours=TOKEN_TTL_HOURS)},
                app.config["SECRET_KEY"],
            )
            return jsonify({"token": token, "role": "admin"})
        
        user = db.query(models.User).filter(models.User.email == email).first()

        if not user or not user.hashed_password or not bcrypt.check_password_hash(user.hashed_password, password):
            app.logger.warning(f"Invalid login attempt for: {email}")
            return jsonify({"message": "Invalid credentials!"}), 401

        token = _issue_token_for(user.email)
        return jsonify({"token": token, "role": getattr(user, "role", "student")})
    finally:
        db.close()

@app.post("/api/google-login")
def google_login():
    db: Session = next(get_db())
    data = request.get_json() or {}
    token_from_frontend = data.get("credential")
    try:
        info = id_token.verify_oauth2_token(
            token_from_frontend,
            g_requests.Request(),
            os.getenv("VITE_GOOGLE_CLIENT_ID"),
        )
        email = info["email"].lower()
        name = info.get("name")

        user = db.query(models.User).filter(models.User.email == email).first()
        if not user:
            user = models.User(name=name, email=email, hashed_password=None)
            db.add(user)
            db.commit()
            db.refresh(user)

        app_token = _issue_token_for(user.email)
        return jsonify({"token": app_token})
    except Exception as e:
        db.rollback()
        app.logger.exception("Error in /api/google-login")
        return jsonify({"message": "Google Sign-In failed.", "error": str(e)}), 401
    finally:
        db.close()

@app.get("/api/me")
@token_required
def me(current_user):
    """Quick check to keep the user logged in after restart."""
    return jsonify({
        "email": current_user.email,
        "name": getattr(current_user, "name", None),
        "role": getattr(current_user, "role", "student"),
    })

@app.post("/api/refresh")
@token_required
def refresh(current_user):
    """Issue a fresh token (same user)."""
    new_token = _issue_token_for(current_user.email)
    return jsonify({"token": new_token})

# =========================
# Document Routes
# =========================
@app.post("/api/documents")
@token_required
def upload_and_process_document(current_user):
    """
    Accepts either:
      - multipart/form-data with 'file' (pdf/docx/txt/png/jpg/jpeg/tiff/bmp/webp)
      - application/json with { "url": "https://..." } to import a web page
      - application/json with { "note": "..." } to save a note
    """
    db: Session = next(get_db())
    try:
        # ---- 1) FILE UPLOAD PATH (multipart/form-data) ----
        if "file" in request.files:
            file = request.files["file"]
            if not file or file.filename.strip() == "":
                return jsonify({"error": "No file selected"}), 400

            filename = secure_filename(file.filename)
            ext = os.path.splitext(filename)[1].lower()
            
            # Validate file extension
            allowed_extensions = ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp']
            if ext not in allowed_extensions:
                return jsonify({"error": f"Unsupported file type: {ext}. Allowed: {', '.join(allowed_extensions)}"}), 400
            
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            
            try:
                file.save(filepath)
                app.logger.info(f"Saved uploaded file to {filepath}")
            except Exception as e:
                app.logger.error(f"Failed to save file: {str(e)}")
                return jsonify({"error": f"Failed to save file: {str(e)}"}), 500

            extracted_text = ""
            
            try:
                if ext == ".pdf":
                    extracted_text = extract_text_from_pdf(filepath)
                    if not extracted_text.strip():
                        app.logger.warning(f"No text extracted from PDF: {filename}")
                        extracted_text = "[PDF contains no extractable text or is image-based]"
                        
                elif ext == ".docx":
                    extracted_text = extract_text_from_docx(filepath)
                    # Create HTML preview
                    try:
                        html_content = convert_docx_to_html(filepath)
                        html_filename = os.path.splitext(filename)[0] + ".html"
                        html_path = os.path.join(app.config["UPLOAD_FOLDER"], html_filename)
                        with open(html_path, "w", encoding="utf-8") as hf:
                            hf.write(html_content)
                        app.logger.info(f"Created HTML preview: {html_path}")
                    except Exception as html_err:
                        app.logger.error(f"Failed to create HTML preview: {str(html_err)}")
                        # Continue anyway - text extraction succeeded
                        
                elif ext == ".txt":
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        extracted_text = f.read()
                        
                elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"]:
                    extracted_text = extract_text_from_image(filepath)
                    if not extracted_text.strip():
                        extracted_text = "[No text detected in image]"
                else:
                    return jsonify({"error": f"Unsupported file type: {ext}"}), 400
                    
            except Exception as extract_err:
                app.logger.error(f"Text extraction failed for {filename}: {str(extract_err)}")
                return jsonify({"error": f"Failed to process file: {str(extract_err)}"}), 500

            # Create document record
            new_document = models.Document(
                filename=filename,
                text=extracted_text,
                user_id=getattr(current_user, "id", None),
            )
            db.add(new_document)
            db.commit()
            db.refresh(new_document)

            # Store in ChromaDB
            try:
                from vector_store import add_document
                add_document(new_document.id, extracted_text)
                app.logger.info(f"Added document {new_document.id} to ChromaDB")
            except Exception as e:
                app.logger.warning(f"ChromaDB add failed (non-critical): {e}")

            return jsonify({
                "message": "Document processed successfully",
                "doc_id": new_document.id,
                "filename": filename,
                "text_length": len(extracted_text)
            }), 200

        # ---- 2) JSON PATH (application/json) ----
        if request.is_json:
            data = request.get_json() or {}
            url = (data.get("url") or "").strip()
            note = data.get("note")

            if url:
                try:
                    parsed = urlparse(url)
                    if not parsed.scheme or not parsed.netloc:
                        return jsonify({"error": "Invalid URL"}), 400
                    text = extract_text_from_url(url)
                    filename = f"web_{parsed.netloc}.txt"
                    new_document = models.Document(
                        filename=filename,
                        text=text,
                        user_id=getattr(current_user, "id", None),
                    )
                    db.add(new_document)
                    db.commit()
                    db.refresh(new_document)
                    return jsonify({"message": "Web page imported", "doc_id": new_document.id}), 200
                except Exception as e:
                    db.rollback()
                    return jsonify({"error": f"Failed to fetch URL: {str(e)}"}), 500

            if isinstance(note, str) and note.strip():
                filename = f"note_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
                new_document = models.Document(
                    filename=filename,
                    text=note.strip(),
                    user_id=getattr(current_user, "id", None),
                )
                db.add(new_document)
                db.commit()
                db.refresh(new_document)
                return jsonify({"message": "Note saved", "doc_id": new_document.id}), 200

            return jsonify({"error": "Provide either 'file', 'url', or 'note'"}), 400

        # ---- 3) No file and not JSON ----
        return jsonify({"error": "No file part and no JSON body"}), 400

    except Exception as e:
        db.rollback()
        app.logger.exception("Error in upload_and_process_document")
        return jsonify({"message": "Error processing document", "error": str(e)}), 500
    finally:
        db.close()

@app.get("/api/documents/<int:doc_id>")
@token_required
def get_document_details(current_user, doc_id: int):
    db: Session = next(get_db())
    try:
        document = db.query(models.Document).filter(models.Document.id == doc_id).first()
        if not document:
            return jsonify({"error": "Document not found"}), 404

        messages = (
            db.query(models.Message)
            .filter(models.Message.document_id == doc_id)
            .order_by(models.Message.created_at.asc())
            .all()
        )

        return jsonify({
            "filename": document.filename,
            "text": document.text,
            "conversation": [
                {
                    "id": m.id,
                    "message_id": m.id if m.role == "assistant" else None,
                    "role": m.role,
                    "content": m.content,
                    "time": m.created_at.strftime("%H:%M") if getattr(m, "created_at", None) else None
                } for m in messages
            ],
        })
    finally:
        db.close()

@app.get("/api/documents")
@token_required
def list_documents(current_user):
    db: Session = next(get_db())
    try:
        q = db.query(models.Document)
        if hasattr(models.Document, "user_id") and getattr(current_user, "id", None):
            q = q.filter(models.Document.user_id == current_user.id)
        docs = q.order_by(models.Document.created_at.desc()).all()
        return jsonify([
            {
                "id": d.id,
                "filename": d.filename,
                "created_at": d.created_at.isoformat() if getattr(d, "created_at", None) else None,
                "text": (d.text or "")[:300],
            } for d in docs
        ])
    finally:
        db.close()

@app.delete("/api/documents/<int:doc_id>")
@token_required
def delete_document(current_user, doc_id: int):
    db: Session = next(get_db())
    try:
        document = db.query(models.Document).filter(models.Document.id == doc_id).first()
        if not document:
            return jsonify({"error": "Document not found"}), 404

        if hasattr(models.Document, "user_id") and getattr(current_user, "id", None):
            if document.user_id != current_user.id:
                return jsonify({"error": "Not authorized to delete this document"}), 403

        try:
            if document.filename:
                path = os.path.join(app.config["UPLOAD_FOLDER"], document.filename)
                if os.path.exists(path):
                    os.remove(path)
        except Exception:
            pass

        db.delete(document)
        db.commit()
        return jsonify({"message": "Document deleted"}), 200
    except Exception as e:
        db.rollback()
        return jsonify({"error": f"Failed to delete document: {str(e)}"}), 500
    finally:
        db.close()

@app.route("/uploads/<path:filename>")
def serve_file(filename: str):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)

# =========================
# Feedback Routes
# =========================
@app.post("/api/feedback")
@token_required
def create_feedback(current_user):
    db: Session = next(get_db())
    data = request.get_json() or {}
    try:
        message_id = data.get("message_id")
        rating = (data.get("rating") or "").lower()
        note = data.get("note")
        comment = data.get("comment")  # Add comment field

        try:
            message_id = int(message_id)
        except (TypeError, ValueError):
            return jsonify({"error": "Invalid message_id"}), 400

        if rating not in ("up", "down"):
            return jsonify({"error": "rating must be 'up' or 'down'"}), 400

        # Ensure message exists
        msg = db.query(models.Message).filter(models.Message.id == message_id).first()
        if not msg:
            return jsonify({"error": "Message not found"}), 404

        # Create feedback with comment
        fb = models.Feedback(
            message_id=message_id, 
            rating=rating, 
            note=note,
            comment=comment  # Store comment
        )
        db.add(fb)
        db.commit()
        db.refresh(fb)

        app.logger.info(f"Feedback saved: {rating} for message {message_id}, comment: {comment[:50] if comment else 'None'}")

        return jsonify({
            "message": "Feedback saved",
            "feedback": {
                "id": fb.id,
                "message_id": fb.message_id,
                "rating": fb.rating,
                "note": fb.note,
                "comment": fb.comment
            }
        }), 200
    except Exception as e:
        db.rollback()
        app.logger.exception("Error saving feedback")
        return jsonify({"error": f"Failed to save feedback: {str(e)}"}), 500
    finally:
        db.close()

@app.get("/api/feedback/<int:message_id>")
@token_required
def get_feedback_summary(current_user, message_id: int):
    db: Session = next(get_db())
    try:
        if not db.query(models.Message.id).filter(models.Message.id == message_id).first():
            return jsonify({"error": "Message not found"}), 404

        up_count = db.query(func.count(models.Feedback.id)).filter(
            models.Feedback.message_id == message_id,
            models.Feedback.rating == "up"
        ).scalar() or 0

        down_count = db.query(func.count(models.Feedback.id)).filter(
            models.Feedback.message_id == message_id,
            models.Feedback.rating == "down"
        ).scalar() or 0

        return jsonify({
            "message_id": message_id,
            "counts": {"up": up_count, "down": down_count}
        }), 200
    finally:
        db.close()

# =========================
# Admin Routes
# =========================
ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "admin@smartdocq.com").lower()
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "Admin@123")

# Add logging to verify credentials are loaded
app.logger.info(f"Admin credentials loaded - Email: {ADMIN_EMAIL}")

def admin_required(f):
    @wraps(f)
    def decorated(current_user, *args, **kwargs):
        user_role = getattr(current_user, "role", "student")
        app.logger.info(f"Admin check - User role: {user_role}")
        if user_role != "admin":
            return jsonify({"message": "Admin access required!"}), 403
        return f(current_user, *args, **kwargs)
    return decorated

@app.get("/api/admin/users")
@token_required
@admin_required
def get_all_users(current_user):
    """Admin-only: Get all users"""
    db: Session = next(get_db())
    try:
        users = db.query(models.User).order_by(models.User.created_at.desc()).all()
        return jsonify([
            {
                "id": u.id,
                "name": u.name,
                "email": u.email,
                "role": getattr(u, "role", "student"),
                "created_at": u.created_at.isoformat() if getattr(u, "created_at", None) else None,
            } for u in users
        ])
    finally:
        db.close()

@app.get("/api/admin/documents")
@token_required
@admin_required
def get_all_documents_admin(current_user):
    """Admin-only: Get all documents from all users"""
    db: Session = next(get_db())
    try:
        docs = db.query(models.Document).order_by(models.Document.created_at.desc()).all()
        return jsonify([
            {
                "id": d.id,
                "filename": d.filename,
                "user_id": d.user_id,
                "created_at": d.created_at.isoformat() if getattr(d, "created_at", None) else None,
                "text_preview": (d.text or "")[:200],
            } for d in docs
        ])
    finally:
        db.close()

@app.get("/api/admin/stats")
@token_required
@admin_required
def get_admin_stats(current_user):
    """Admin-only: Get system statistics with analytics data"""
    db: Session = next(get_db())
    try:
        total_users = db.query(func.count(models.User.id)).scalar() or 0
        total_docs = db.query(func.count(models.Document.id)).scalar() or 0
        total_messages = db.query(func.count(models.Message.id)).scalar() or 0
        total_feedback = db.query(func.count(models.Feedback.id)).scalar() or 0
        
        # Feedback breakdown
        positive_feedback = db.query(func.count(models.Feedback.id)).filter(
            models.Feedback.rating == "up"
        ).scalar() or 0
        negative_feedback = db.query(func.count(models.Feedback.id)).filter(
            models.Feedback.rating == "down"
        ).scalar() or 0
        
        # User roles breakdown
        admin_users = db.query(func.count(models.User.id)).filter(
            models.User.role == "admin"
        ).scalar() or 0
        student_users = total_users - admin_users
        
        # Document types breakdown (count by file extension)
        from sqlalchemy import case
        doc_types = db.query(
            case(
                (models.Document.filename.like('%.pdf'), 'PDF'),
                (models.Document.filename.like('%.docx'), 'Word'),
                (models.Document.filename.like('%.txt'), 'Text'),
                (models.Document.filename.like('%.png'), 'Image'),
                (models.Document.filename.like('%.jpg'), 'Image'),
                (models.Document.filename.like('%.jpeg'), 'Image'),
                else_='Other'
            ).label('type'),
            func.count(models.Document.id).label('count')
        ).group_by('type').all()
        
        doc_type_stats = {dtype: count for dtype, count in doc_types}
        
        # Activity over last 7 days
        from datetime import datetime, timedelta
        today = datetime.now().date()
        activity_data = []
        for i in range(6, -1, -1):
            date = today - timedelta(days=i)
            date_start = datetime.combine(date, datetime.min.time())
            date_end = datetime.combine(date, datetime.max.time())
            
            docs_count = db.query(func.count(models.Document.id)).filter(
                models.Document.created_at >= date_start,
                models.Document.created_at <= date_end
            ).scalar() or 0
            
            messages_count = db.query(func.count(models.Message.id)).filter(
                models.Message.created_at >= date_start,
                models.Message.created_at <= date_end
            ).scalar() or 0
            
            activity_data.append({
                "date": date.strftime("%b %d"),
                "documents": docs_count,
                "messages": messages_count
            })
        
        return jsonify({
            "total_users": total_users,
            "total_documents": total_docs,
            "total_messages": total_messages,
            "total_feedback": total_feedback,
            "feedback_breakdown": {
                "positive": positive_feedback,
                "negative": negative_feedback
            },
            "user_roles": {
                "admin": admin_users,
                "student": student_users
            },
            "document_types": doc_type_stats,
            "activity_timeline": activity_data
        })
    finally:
        db.close()

@app.delete("/api/admin/users/<int:user_id>")
@token_required
@admin_required
def delete_user_admin(current_user, user_id: int):
    """Admin-only: Delete a user and all their documents"""
    db: Session = next(get_db())
    try:
        user = db.query(models.User).filter(models.User.id == user_id).first()
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        # Prevent deleting admin users
        if getattr(user, "role", "student") == "admin":
            return jsonify({"error": "Cannot delete admin users"}), 403
        
        # Delete user's documents first (cascade should handle this, but explicit is better)
        user_docs = db.query(models.Document).filter(models.Document.user_id == user_id).all()
        for doc in user_docs:
            try:
                # Delete physical file if exists
                if doc.filename:
                    path = os.path.join(app.config["UPLOAD_FOLDER"], doc.filename)
                    if os.path.exists(path):
                        os.remove(path)
            except Exception as e:
                app.logger.warning(f"Failed to delete file for doc {doc.id}: {e}")
            db.delete(doc)
        
        # Delete the user
        db.delete(user)
        db.commit()
        
        return jsonify({"message": f"User {user.email} and their documents deleted successfully"}), 200
    except Exception as e:
        db.rollback()
        app.logger.exception("Error deleting user")
        return jsonify({"error": f"Failed to delete user: {str(e)}"}), 500
    finally:
        db.close()

@app.get("/api/admin/feedbacks")
@token_required
@admin_required
def get_all_feedbacks(current_user):
    """Admin-only: Get all feedbacks with notes"""
    db: Session = next(get_db())
    try:
        feedbacks = db.query(models.Feedback).order_by(models.Feedback.created_at.desc()).all()
        return jsonify([
            {
                "id": f.id,
                "message_id": f.message_id,
                "rating": f.rating,
                "note": f.note,
                "created_at": f.created_at.isoformat() if getattr(f, "created_at", None) else None,
            } for f in feedbacks
        ])
    finally:
        db.close()

@app.get("/api/admin/feedback")
@token_required
@admin_required
def get_all_feedback_admin(current_user):
    """Admin-only: Get all feedback including comments"""
    db: Session = next(get_db())
    try:
        # Get feedback with message content
        feedbacks = db.query(
            models.Feedback,
            models.Message
        ).join(
            models.Message,
            models.Feedback.message_id == models.Message.id
        ).order_by(
            models.Feedback.created_at.desc()
        ).all()
        
        return jsonify([
            {
                "id": fb.id,
                "message_id": fb.message_id,
                "rating": fb.rating,
                "comment": fb.comment,
                "note": fb.note,
                "created_at": fb.created_at.isoformat() if fb.created_at else None,
                "message_content": msg.content[:200] if msg else None,  # First 200 chars
                "document_id": msg.document_id if msg else None
            } for fb, msg in feedbacks
        ])
    finally:
        db.close()

@app.get("/api/admin/feedback/negative")
@token_required
@admin_required
def get_negative_feedback_admin(current_user):
    """Admin-only: Get all negative feedback with comments"""
    db: Session = next(get_db())
    try:
        feedbacks = db.query(
            models.Feedback,
            models.Message
        ).join(
            models.Message,
            models.Feedback.message_id == models.Message.id
        ).filter(
            models.Feedback.rating == 'down'
        ).order_by(
            models.Feedback.created_at.desc()
        ).all()
        
        return jsonify([
            {
                "id": fb.id,
                "message_id": fb.message_id,
                "comment": fb.comment,
                "created_at": fb.created_at.isoformat() if fb.created_at else None,
                "message_content": msg.content if msg else None,
                "document_id": msg.document_id if msg else None
            } for fb, msg in feedbacks if fb.comment  # Only with comments
        ])
    finally:
        db.close()

# =========================
# Q&A Route (Gemini guarded)
# =========================
@app.post("/api/ask")
@token_required
def ask_question(current_user):
    db: Session = next(get_db())
    data = request.get_json() or {}

    raw_doc_id = data.get("doc_id")
    try:
        doc_id = int(raw_doc_id)
    except (TypeError, ValueError):
        return jsonify({"error": "Invalid doc_id; must be an integer."}), 400

    question = (data.get("question") or "").strip()
    if not question:
        return jsonify({"error": "Empty question"}), 400

    force_mock = request.headers.get("x-ai-mock", "").lower() == "true"

    try:
        user_msg = models.Message(role="user", content=question, document_id=doc_id)
        db.add(user_msg)
        db.commit()

        document = db.query(models.Document).filter(models.Document.id == doc_id).first()
        if not document:
            return jsonify({"error": "Document context not found"}), 404

        if force_mock or (not AI_ENABLED) or (genai is None):
            payload = safe_answer_mock("AI disabled or not initialized")
            assistant_msg = models.Message(
                role="assistant", content=payload["answer"], document_id=doc_id
            )
            db.add(assistant_msg)
            db.commit()
            db.refresh(assistant_msg)
            payload["message_id"] = assistant_msg.id
            return jsonify(payload), 200

        # ---- Live Gemini AI ----
        from prompts import build_prompt, parse_llm_json
        # Retrieve relevant text from ChromaDB to enrich context
        try:
            from vector_store import query_similar
            similar_chunks, _ = query_similar(question)
            context_from_chroma = " ".join(similar_chunks)
        except Exception as e:
            app.logger.warning(f"ChromaDB query failed: {e}")
            context_from_chroma = ""

        # Use both DB text and retrieved context
        safe_context = ((document.text or "")[:3000] + "\n" + context_from_chroma)[:6000]

        prompt = build_prompt(context=safe_context, question=question)
        model = genai.GenerativeModel(MODEL_NAME)

        try:
            response = model.generate_content(prompt, generation_config={"temperature": TEMPERATURE})
        except Exception:
            if TESTING_MODE_ON_RATE_LIMIT:
                payload = {"answer": "This is a testing response (rate limit).",
                           "_used_model": MODEL_NAME, "_mode": "rate-limit-mock"}
                assistant_msg = models.Message(role="assistant", content=payload["answer"], document_id=doc_id)
                db.add(assistant_msg)
                db.commit()
                db.refresh(assistant_msg)
                payload["message_id"] = assistant_msg.id
                return jsonify(payload), 200
            raise

        response_text = getattr(response, "text", "") or ""
        try:
            parsed = parse_llm_json(response_text)
        except Exception:
            parsed = {"answer": response_text}

        assistant_msg = models.Message(
            role="assistant", content=parsed.get("answer") or "", document_id=doc_id
        )
        db.add(assistant_msg)
        db.commit()
        db.refresh(assistant_msg)

        parsed["_used_model"] = MODEL_NAME
        parsed["_mode"] = "live"
        parsed["message_id"] = assistant_msg.id
        return jsonify(parsed), 200

    except Exception:
        import traceback
        db.rollback()
        fail_text = "Failed to get response: " + traceback.format_exc()
        assistant_msg = models.Message(role="assistant", content=fail_text, document_id=doc_id)
        db.add(assistant_msg)
        db.commit()
        db.refresh(assistant_msg)
        return jsonify({"error": fail_text, "message_id": assistant_msg.id}), 500
    finally:
        db.close()

@app.post("/api/general-ask")
@token_required
def general_ask(current_user):
    """
    General Q&A endpoint - answers any question without document context.
    Uses the AI model directly for open-ended queries.
    """
    db: Session = next(get_db())
    data = request.get_json() or {}
    question = (data.get("question") or "").strip()
    
    if not question:
        return jsonify({"error": "Empty question"}), 400

    force_mock = request.headers.get("x-ai-mock", "").lower() == "true"

    try:
        # Store user message (optional - for logging)
        user_msg = models.Message(role="user", content=question, document_id=None)
        db.add(user_msg)
        db.commit()

        if force_mock or (not AI_ENABLED) or (genai is None):
            answer_text = f"(mock) AI disabled. Your question was: {question}"
            assistant_msg = models.Message(role="assistant", content=answer_text, document_id=None)
            db.add(assistant_msg)
            db.commit()
            return jsonify({"answer": answer_text, "_mode": "mock"}), 200

        # Simple prompt for general knowledge
        prompt = f"""You are a helpful AI assistant. Answer the following question clearly and concisely.
If you don't know the answer, say so honestly.

Question: {question}

Answer:"""

        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(prompt, generation_config={"temperature": TEMPERATURE})
        
        answer = getattr(response, "text", "") or "No response generated."
        
        # Store assistant response
        assistant_msg = models.Message(role="assistant", content=answer, document_id=None)
        db.add(assistant_msg)
        db.commit()
        
        return jsonify({
            "answer": answer,
            "_used_model": MODEL_NAME,
            "_mode": "live"
        }), 200

    except Exception as e:
        db.rollback()
        app.logger.exception("Error in /api/general-ask")
        if TESTING_MODE_ON_RATE_LIMIT:
            return jsonify({
                "answer": "This is a testing response (rate limit reached).",
                "_mode": "rate-limit-mock"
            }), 200
        return jsonify({"error": f"Failed to get response: {str(e)}"}), 500

# =========================
# Share Routes
# =========================
@app.post("/api/share")
@token_required
def create_share_link(current_user):
    """Create a shareable link for a document conversation"""
    db: Session = next(get_db())
    data = request.get_json() or {}
    
    try:
        doc_id = int(data.get("doc_id"))
    except (TypeError, ValueError):
        return jsonify({"error": "Invalid doc_id"}), 400
    
    try:
        # Verify document exists and user has access
        document = db.query(models.Document).filter(models.Document.id == doc_id).first()
        if not document:
            return jsonify({"error": "Document not found"}), 404
        
        if hasattr(models.Document, "user_id") and getattr(current_user, "id", None):
            if document.user_id != current_user.id:
                return jsonify({"error": "Not authorized to share this document"}), 403
        
        # Check if share already exists
        existing_share = db.query(models.SharedConversation).filter(
            models.SharedConversation.document_id == doc_id,
            models.SharedConversation.is_active == 1
        ).first()
        
        if existing_share:
            share_id = existing_share.share_id
        else:
            # Create new share
            import secrets
            share_id = secrets.token_urlsafe(16)
            
            # Optional: set expiration (e.g., 30 days from now)
            # expires_at = datetime.now(timezone.utc) + timedelta(days=30)
            
            new_share = models.SharedConversation(
                share_id=share_id,
                document_id=doc_id,
                created_by=getattr(current_user, "id", None),
                # expires_at=expires_at,
            )
            db.add(new_share)
            db.commit()
            db.refresh(new_share)
        
        # Build full URL
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:5173")
        share_url = f"{frontend_url}/share/{share_id}"
        
        return jsonify({
            "message": "Share link created",
            "share_id": share_id,
            "share_url": share_url
        }), 200
        
    except Exception as e:
        db.rollback()
        app.logger.exception("Error creating share link")
        return jsonify({"error": f"Failed to create share link: {str(e)}"}), 500
    finally:
        db.close()

@app.get("/api/share/<share_id>")
def get_shared_conversation(share_id: str):
    """Public endpoint to retrieve shared conversation - no auth required"""
    db: Session = next(get_db())
    try:
        # Find share record
        share = db.query(models.SharedConversation).filter(
            models.SharedConversation.share_id == share_id,
            models.SharedConversation.is_active == 1
        ).first()
        
        if not share:
            return jsonify({"error": "Share not found or has been disabled"}), 404
        
        # Check if expired
        if share.expires_at and datetime.now(timezone.utc) > share.expires_at:
            return jsonify({"error": "Share link has expired"}), 410
        
        # Get document
        document = db.query(models.Document).filter(
            models.Document.id == share.document_id
        ).first()
        
        if not document:
            return jsonify({"error": "Document not found"}), 404
        
        # Get messages
        messages = (
            db.query(models.Message)
            .filter(models.Message.document_id == share.document_id)
            .order_by(models.Message.created_at.asc())
            .all()
        )
        
        return jsonify({
            "filename": document.filename,
            "text": document.text,
            "conversation": [
                {
                    "role": m.role,
                    "content": m.content,
                    "time": m.created_at.strftime("%H:%M") if getattr(m, "created_at", None) else None
                } for m in messages
            ],
            "created_at": document.created_at.isoformat() if getattr(document, "created_at", None) else None,
        })
    finally:
        db.close()

@app.delete("/api/share/<share_id>")
@token_required
def delete_share_link(current_user, share_id: str):
    """Disable a share link"""
    db: Session = next(get_db())
    try:
        share = db.query(models.SharedConversation).filter(
            models.SharedConversation.share_id == share_id
        ).first()
        
        if not share:
            return jsonify({"error": "Share not found"}), 404
        
        # Verify ownership
        if getattr(current_user, "id", None) and share.created_by != current_user.id:
            return jsonify({"error": "Not authorized to delete this share"}), 403
        
        share.is_active = 0
        db.commit()
        
        return jsonify({"message": "Share link disabled"}), 200
    except Exception as e:
        db.rollback()
        return jsonify({"error": f"Failed to disable share: {str(e)}"}), 500
    finally:
        db.close()

# =========================
# Health Check Routes
# =========================
@app.route("/")
def index():
    return jsonify({"message": "SmartDocQ API is running"}), 200

@app.route("/health", methods=['GET'])
def health_check():
    """Health check endpoint to prevent service from sleeping"""
    return jsonify({
        "status": "healthy",
        "timestamp": datetime.now().isoformat()
    }), 200

@app.get("/api/health")
def api_health():
    """Detailed health check"""
    return jsonify({
        "status": "healthy",
        "database": "connected",
        "ai_model": MODEL_NAME if AI_ENABLED else "disabled",
        "frontend_url": os.getenv("FRONTEND_URL", "not set")
    }), 200

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5001))  # Changed to 5001 (5000 is AirPlay on macOS)
    print(f"🚀 Starting server on http://localhost:{port}")
    app.run(host="0.0.0.0", port=port, debug=True)